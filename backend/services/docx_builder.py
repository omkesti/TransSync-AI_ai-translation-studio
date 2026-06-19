"""
docx_builder.py
───────────────
Shared DOCX reconstruction logic.

Used by:
    - backend/routes/export.py       (single-doc export)
    - backend/routes/export_batch.py  (multi-doc ZIP export)

Responsibility:
    Given raw_text (paragraph-separated by \\n\\n) and a list of
    { source, translation } items, rebuild a clean DOCX that mirrors
    the original document structure with translated content in place.
"""

import base64
import io
import re
from datetime import datetime
from typing import Any, Dict, Iterator, List, Optional

from pydantic import BaseModel
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from backend.services.document_parser import inject_translation_into_paragraph


# ── Shared schemas ────────────────────────────────────────────────────────────

class TranslationItem(BaseModel):
    source:      str
    translation: str
    match_type:  Optional[str] = None
    # Optional unit id for ID-based reconstruction (unused by the OOXML path,
    # which matches by paragraph text). Kept for forward compatibility.
    id:          Optional[str] = None


class ExtractionUnitSchema(BaseModel):
    """
    One addressable unit of a parsed document (a paragraph or table cell).
    Carries the text plus enough positional/formatting context to align a
    translation back onto the original node. Currently informational — the
    OOXML injection path matches by paragraph text, not by unit id.
    """
    text:       str
    formatting: Dict[str, Any] = {}
    position:   Dict[str, Any] = {}


class DocExportData(BaseModel):
    """
    All the data needed to reconstruct one translated document.
    Both the single-export and batch-export routes build instances of this.

    Two export strategies are supported, selected by which fields are present:
      - original_docx_b64 set  → OOXML run-level injection (format-preserving).
                                  Used when the source upload was a .docx.
      - original_docx_b64 unset → legacy from-scratch reconstruction (raw_text).
                                  Used when the source upload was a PDF.
    """
    doc_id:      str
    filename:    str = "document"
    source_lang: str = "en"
    target_lang: str
    raw_text:    str = ""
    translations: List[TranslationItem]
    # Base64 of the ORIGINAL .docx bytes — enables format-preserving export.
    original_docx_b64: Optional[str] = None
    # Optional structured extraction payload (reserved for ID-based path).
    extraction_data:   Optional[dict] = None


# ── XML safety ────────────────────────────────────────────────────────────────

# OOXML (XML 1.0) forbids C0 control characters except tab (\x09), newline
# (\x0A) and carriage return (\x0D). PDF text extraction and, rarely, LLM output
# can carry NULL bytes / form feeds / other control chars; passing them to
# python-docx raises "All strings must be XML compatible". Strip them defensively
# at every text-injection point.
_XML_ILLEGAL_RE = re.compile(r"[\x00-\x08\x0B\x0C\x0E-\x1F]")


def xml_safe(text: str) -> str:
    """Remove XML-incompatible control characters from text (keeps \\t \\n \\r)."""
    if not text:
        return text
    return _XML_ILLEGAL_RE.sub("", text)


# ── Heading detection heuristic ───────────────────────────────────────────────

def is_heading(text: str) -> bool:
    """
    Heuristically detect heading-like paragraphs.
    A paragraph is treated as a heading if it:
      - Is short (≤ 80 chars)
      - Does not end with sentence-ending punctuation
      - Is not just a number
    """
    text = text.strip()
    if not text or len(text) > 80:
        return False
    if text[-1] in ".!?,;:":
        return False
    if re.match(r"^\d+\.?$", text):
        return False
    return True


# ── Paragraph reconstruction ─────────────────────────────────────────────────

def build_fuzzy_pattern(source: str) -> re.Pattern:
    """
    Builds a case-insensitive regex that matches the source sentence
    while tolerating whitespace variations (e.g., extra spaces or line breaks).
    """
    words = source.strip().split()
    if not words:
        return re.compile(re.escape(source.strip()), flags=re.IGNORECASE)
    pattern_str = r"\s+".join(re.escape(word) for word in words)
    return re.compile(pattern_str, flags=re.IGNORECASE)


def reconstruct_paragraphs(
    raw_text: str,
    translations: List[TranslationItem],
) -> List[str]:
    """
    Splits raw_text into paragraphs (by double newline) and replaces each
    source sentence with its translation using fuzzy whitespace matching.

    Returns a list of translated paragraph strings in original order.
    Falls back to flat sentence list if raw_text is empty.
    """
    lookup = {item.source.strip(): item.translation for item in translations}

    if not raw_text.strip():
        return [item.translation for item in translations]

    paragraphs = [p.strip() for p in re.split(r"\n{2,}", raw_text) if p.strip()]

    result = []
    for paragraph in paragraphs:
        translated = paragraph
        for source, translation in lookup.items():
            if not source:
                continue
            pattern = build_fuzzy_pattern(source)
            translated = pattern.sub(translation, translated, count=1)
        result.append(translated)

    return result


# ── DOCX builder ──────────────────────────────────────────────────────────────

def add_horizontal_rule(doc: Document) -> None:
    """Inserts a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def reconstruct_docx(data: DocExportData) -> io.BytesIO:
    """
    Builds a complete output DOCX from a DocExportData object:
        1. Translated paragraphs in original document order
        2. A thin horizontal rule separator
        3. One-line italic attribution footer
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # Reconstruct paragraphs
    translated_paragraphs = reconstruct_paragraphs(data.raw_text, data.translations)

    for text in translated_paragraphs:
        text = xml_safe(text)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(11)

        if is_heading(text):
            run.bold = True
            run.font.size = Pt(13)

    # Footer separator
    doc.add_paragraph()
    add_horizontal_rule(doc)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Translated by TransSync AI  ·  "
        f"{data.source_lang.upper()} → {data.target_lang.upper()}  ·  "
        f"{datetime.utcnow().strftime('%B %d, %Y')}"
    )
    footer_run.italic    = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Serialize
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def make_output_filename(filename: str, target_lang: str) -> str:
    """Derives a clean output filename like translated_report_fr.docx"""
    base = filename.rsplit(".", 1)[0] if "." in filename else filename
    return f"translated_{base}_{target_lang}.docx"


# ── OOXML format-preserving export (DOCX path) ────────────────────────────────
#
# build_translated_docx() opens the ORIGINAL .docx and injects translations
# directly into existing runs via inject_translation_into_paragraph(). Nothing
# is rebuilt, so tables, images, styles, headers/footers and inline formatting
# all survive untouched. Paragraphs with no matching translation are left as-is.


def _iter_all_paragraphs(doc: Document) -> Iterator[Paragraph]:
    """
    Yield every paragraph in document order: body paragraphs first, then every
    paragraph inside every table cell, recursing through nested tables.
    """
    yield from doc.paragraphs

    def _recurse_table(table: Table) -> Iterator[Paragraph]:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para
                for nested in cell.tables:
                    yield from _recurse_table(nested)

    for table in doc.tables:
        yield from _recurse_table(table)


def _resolve_paragraph_translation(
    text: str,
    exact_map: Dict[str, str],
    fuzzy_items: List[tuple],
) -> Optional[str]:
    """
    Compute the translated text for one original paragraph.

    1. Whole-paragraph exact match (the common case — one paragraph is one
       translation unit).
    2. Otherwise substitute each source sentence with its translation inside
       the paragraph, tolerating whitespace differences. This handles
       paragraphs the NLP layer split into several sentences.

    Returns the new text, or None if nothing matched (caller leaves the
    paragraph untouched, preserving it perfectly).
    """
    key = text.strip()
    if key in exact_map:
        return exact_map[key]

    result = text
    changed = False
    for source, translation in fuzzy_items:
        pattern = build_fuzzy_pattern(source)
        # Replace via a function so backslashes/group refs in the translation
        # are treated as literal text, not regex replacement syntax.
        new_result = pattern.sub(lambda _m: translation, result, count=1)
        if new_result != result:
            result = new_result
            changed = True

    return result if changed else None


def build_translated_docx(data: DocExportData) -> io.BytesIO:
    """
    Format-preserving DOCX export via OOXML run-level injection.

    Opens the original document from `data.original_docx_b64`, walks every
    paragraph (body + table cells, nested tables included), and injects the
    matching translation into each paragraph's runs while preserving all
    formatting. Returns a BytesIO positioned at 0.

    Raises ValueError if `original_docx_b64` is missing or not decodable.
    """
    if not data.original_docx_b64:
        raise ValueError(
            "build_translated_docx requires original_docx_b64 "
            "(the base64-encoded original .docx)."
        )

    try:
        raw_bytes = base64.b64decode(data.original_docx_b64)
    except Exception as e:  # noqa: BLE001 — surface as a clean ValueError
        raise ValueError(f"original_docx_b64 is not valid base64: {e}")

    doc = Document(io.BytesIO(raw_bytes))

    # Build lookup tables once. exact_map keeps the FIRST translation seen for a
    # given source so duplicate sources stay deterministic.
    exact_map: Dict[str, str] = {}
    fuzzy_items: List[tuple] = []
    for item in data.translations:
        source = (item.source or "").strip()
        if not source or not item.translation:
            continue
        exact_map.setdefault(source, item.translation)
        fuzzy_items.append((source, item.translation))

    for para in _iter_all_paragraphs(doc):
        original = para.text
        if not original.strip():
            continue
        translated = _resolve_paragraph_translation(original, exact_map, fuzzy_items)
        if translated is not None and translated != original:
            inject_translation_into_paragraph(para, xml_safe(translated))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
