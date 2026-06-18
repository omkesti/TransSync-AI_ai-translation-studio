"""
document_parser.py
------------------
Extracts raw plain text from uploaded PDF or DOCX files.

Dependencies:
    pip install PyMuPDF python-docx

Called by:
    backend/routes/upload.py  →  parse_document(filepath)
"""

import os
import fitz          # PyMuPDF — for PDF
import docx          # python-docx — for DOCX


def parse_document(filepath: str) -> str:
    """
    Accepts an absolute file path to a PDF or DOCX.
    Returns the extracted text as a single string.
    Raises ValueError if the file type is unsupported.
    """
    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return _parse_pdf(filepath)
    elif ext == ".docx":
        return _parse_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: '{ext}'. Only .pdf and .docx are accepted.")


def _parse_pdf(filepath: str) -> str:
    """
    Uses PyMuPDF (fitz) to extract text page by page.
    Joins pages with double newlines to preserve structure.
    """
    text_parts = []

    with fitz.open(filepath) as doc:
        for page in doc:
            page_text = page.get_text("text")   # plain text, no HTML
            if page_text.strip():
                text_parts.append(page_text.strip())

    if not text_parts:
        raise ValueError("PDF appears to be empty or is a scanned image (no extractable text).")

    return "\n\n".join(text_parts)


def _parse_docx(filepath: str) -> str:
    """
    Uses python-docx to extract text from the document body AND from every table
    cell (nested tables included). Skips empty paragraphs.

    Paragraphs are joined with blank lines (\\n\\n) so the NLP cleaner treats
    each one as a hard boundary. This mirrors the traversal order used by the
    OOXML injection engine (_iter_all_paragraphs in docx_builder), so every
    extracted unit can be matched back onto its original node at export time.
    """
    doc = docx.Document(filepath)

    parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        _collect_table_text(table, parts)

    if not parts:
        raise ValueError("DOCX file appears to be empty.")

    return "\n\n".join(parts)


def _collect_table_text(table, out: list) -> None:
    """
    Depth-first collection of non-empty paragraph texts from a table, recursing
    through nested tables: table -> row -> cell -> paragraph. Order matches
    docx_builder._iter_all_paragraphs so extraction and injection stay aligned.
    """
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = para.text.strip()
                if text:
                    out.append(text)
            for nested in cell.tables:
                _collect_table_text(nested, out)


# ── OOXML run-level injection ─────────────────────────────────────────────────
#
# WHY THIS EXISTS
# ---------------
# python-docx's `paragraph.text = "..."` destroys ALL Run (<w:r>) nodes inside
# the paragraph, permanently deleting their formatting properties (<w:rPr>:
# bold/italic/font/size/color). To translate a .docx WITHOUT degrading its
# layout we must operate on the ORIGINAL runs instead of rebuilding the file.
#
# A paragraph (<w:p>) holds one or more Runs (<w:r>). Each Run carries a text
# node (<w:t>) and optional formatting (<w:rPr>). The correct injection is:
#
#     runs[0].text = translated_text        # whole translation goes here
#     for run in runs[1:]: run.text = ""     # blank — NOT removed
#
# Blanking (rather than deleting) the trailing runs keeps their XML nodes — and
# therefore the document's structural integrity — fully intact, while the
# translated text safely inherits the paragraph's primary run formatting.

def inject_translation_into_paragraph(paragraph, translated_text: str) -> None:
    """
    Replace the visible text of a paragraph with `translated_text` while
    preserving every Run node and its <w:rPr> formatting.

    - runs[0] receives the full translated string.
    - runs[1:] are blanked (text set to "") but kept in the XML tree.
    - If the paragraph has no runs (e.g. a bare <w:p>), a single run is added
      as a fallback so the text is not lost.

    Mutates the paragraph in place; returns None.
    """
    runs = paragraph.runs

    if not runs:
        # No run to inherit formatting from — append a plain run.
        paragraph.add_run(translated_text)
        return

    runs[0].text = translated_text
    for run in runs[1:]:
        run.text = ""
