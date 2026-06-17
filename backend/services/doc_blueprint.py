"""
doc_blueprint.py
----------------
DOCX-only extraction, persistence, and reconstruction helpers.

This module introduces an ID-based blueprint for addressable DOCX units so
translation and reinsertion do not depend on positional alignment.
"""

from __future__ import annotations

import io
import json
import os
import re
import shutil
import tempfile
import uuid
from datetime import datetime, timezone
from typing import Iterable, Optional

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.section import _Header, _Footer
from docx.shared import Pt, RGBColor
from docx.table import Table, _Cell, _Row
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from pydantic import BaseModel, Field


DOC_BLUEPRINT_ROOT = os.path.join(tempfile.gettempdir(), "transsync_doc_blueprints")
DOC_BLUEPRINT_TTL_SECONDS = 24 * 60 * 60


class RunSpan(BaseModel):
    start: int
    end: int
    text: str = ""
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    strike: Optional[bool] = None
    all_caps: Optional[bool] = None
    small_caps: Optional[bool] = None
    font_name: Optional[str] = None
    font_size_pt: Optional[float] = None
    color_rgb: Optional[str] = None
    highlight_color: Optional[str] = None
    style_name: Optional[str] = None


class DocUnit(BaseModel):
    unit_id: str
    story_type: str
    section_index: int
    container_path: list[str] = Field(default_factory=list)
    paragraph_index: int
    source_text: str
    paragraph_style: str = ""
    run_spans: list[RunSpan] = Field(default_factory=list)
    is_heading: bool = False


class SegmentRef(BaseModel):
    segment_id: str
    unit_id: str
    segment_index_within_unit: int
    text: str
    unit_kind: str = "paragraph"


class DocBlueprint(BaseModel):
    doc_id: str
    source_type: str = "docx"
    reconstruction_mode: str = "blueprint"
    original_filename: str
    source_path: str
    extracted_at: str
    raw_text: str
    units: list[DocUnit] = Field(default_factory=list)


class TranslationItem(BaseModel):
    source: str
    translation: str
    match_type: Optional[str] = None
    score: Optional[float] = None
    segment_id: Optional[str] = None
    unit_id: Optional[str] = None
    segment_index_within_unit: Optional[int] = None


class DocExportData(BaseModel):
    doc_id: str
    filename: str = "document"
    source_lang: str = "en"
    target_lang: str
    raw_text: str = ""
    source_type: str = "docx"
    translations: list[TranslationItem]


def ensure_blueprint_root() -> None:
    os.makedirs(DOC_BLUEPRINT_ROOT, exist_ok=True)


def cleanup_expired_workspaces(ttl_seconds: int = DOC_BLUEPRINT_TTL_SECONDS) -> None:
    ensure_blueprint_root()
    now = datetime.now(timezone.utc).timestamp()
    for name in os.listdir(DOC_BLUEPRINT_ROOT):
        workspace = os.path.join(DOC_BLUEPRINT_ROOT, name)
        if not os.path.isdir(workspace):
            continue
        try:
            modified = os.path.getmtime(workspace)
        except OSError:
            continue
        if now - modified > ttl_seconds:
            shutil.rmtree(workspace, ignore_errors=True)


def get_workspace_path(doc_id: str) -> str:
    ensure_blueprint_root()
    return os.path.join(DOC_BLUEPRINT_ROOT, doc_id)


def get_blueprint_path(doc_id: str) -> str:
    return os.path.join(get_workspace_path(doc_id), "blueprint.json")


def get_source_doc_path(doc_id: str) -> Optional[str]:
    workspace = get_workspace_path(doc_id)
    if not os.path.isdir(workspace):
        return None
    for name in os.listdir(workspace):
        if name.startswith("source."):
            return os.path.join(workspace, name)
    return None


def save_docx_source(doc_id: str, ext: str, source_path: str) -> str:
    workspace = get_workspace_path(doc_id)
    os.makedirs(workspace, exist_ok=True)
    target = os.path.join(workspace, f"source{ext}")
    shutil.move(source_path, target)
    return target


def save_blueprint(blueprint: DocBlueprint) -> None:
    path = get_blueprint_path(blueprint.doc_id)
    os.makedirs(os.path.dirname(path), exist_ok=True)
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(blueprint.model_dump_json(indent=2))


def load_blueprint(doc_id: str) -> Optional[DocBlueprint]:
    path = get_blueprint_path(doc_id)
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as fh:
        return DocBlueprint.model_validate_json(fh.read())


def build_docx_blueprint(doc_id: str, filename: str, source_path: str) -> DocBlueprint:
    doc = Document(source_path)
    units: list[DocUnit] = []
    raw_text_parts: list[str] = []

    units.extend(_extract_story_units(doc, "body", 0))

    for section_index, section in enumerate(doc.sections):
        units.extend(_extract_story_units(section.header, "header", section_index))
        units.extend(_extract_story_units(section.footer, "footer", section_index))

    for unit in units:
        if unit.source_text.strip():
            raw_text_parts.append(unit.source_text.strip())

    blueprint = DocBlueprint(
        doc_id=doc_id,
        original_filename=filename,
        source_path=source_path,
        extracted_at=datetime.now(timezone.utc).isoformat(),
        raw_text="\n\n".join(raw_text_parts),
        units=units,
    )
    return blueprint


def _extract_story_units(root, story_type: str, section_index: int) -> list[DocUnit]:
    units: list[DocUnit] = []
    _walk_container(
        parent=root,
        story_type=story_type,
        section_index=section_index,
        container_path=[],
        out=units,
    )
    return units


def _walk_container(parent, story_type: str, section_index: int, container_path: list[str], out: list[DocUnit]) -> None:
    paragraph_index = 0
    table_index = 0

    for block_type, block in iter_block_items(parent):
        if block_type == "paragraph":
            paragraph = block
            text = paragraph.text or ""
            if text.strip():
                run_spans = _extract_run_spans(paragraph)
                unit_id = _make_unit_id(story_type, section_index, container_path, paragraph_index)
                out.append(
                    DocUnit(
                        unit_id=unit_id,
                        story_type=story_type,
                        section_index=section_index,
                        container_path=list(container_path),
                        paragraph_index=paragraph_index,
                        source_text=text,
                        paragraph_style=getattr(getattr(paragraph, "style", None), "name", "") or "",
                        run_spans=run_spans,
                        is_heading=_is_heading_paragraph(paragraph),
                    )
                )
            paragraph_index += 1
            continue

        table = block
        current_table_index = table_index
        table_index += 1
        table_token = f"table:{current_table_index}"

        for row_index, row in enumerate(table.rows):
            row_token = f"row:{row_index}"
            for cell_index, cell in enumerate(row.cells):
                cell_token = f"cell:{cell_index}"
                _walk_container(
                    parent=cell,
                    story_type=story_type,
                    section_index=section_index,
                    container_path=[*container_path, table_token, row_token, cell_token],
                    out=out,
                )


def iter_block_items(parent) -> Iterable[tuple[str, object]]:
    parent_elm = None
    if isinstance(parent, DocxDocument):
        parent_elm = parent.element.body
    elif isinstance(parent, (_Header, _Footer)):
        parent_elm = parent._element
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    elif isinstance(parent, _Row):
        parent_elm = parent._tr
    else:
        parent_elm = getattr(parent, "_element", None)

    if parent_elm is None:
        return

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield "paragraph", Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield "table", Table(child, parent)


def _extract_run_spans(paragraph: Paragraph) -> list[RunSpan]:
    spans: list[RunSpan] = []
    cursor = 0
    for run in paragraph.runs:
        text = run.text or ""
        end = cursor + len(text)
        spans.append(
            RunSpan(
                start=cursor,
                end=end,
                text=text,
                bold=run.bold,
                italic=run.italic,
                underline=run.underline,
                strike=run.font.strike,
                all_caps=run.font.all_caps,
                small_caps=run.font.small_caps,
                font_name=run.font.name,
                font_size_pt=float(run.font.size.pt) if run.font.size is not None else None,
                color_rgb=str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
                highlight_color=str(run.font.highlight_color) if run.font.highlight_color else None,
                style_name=getattr(getattr(run, "style", None), "name", None),
            )
        )
        cursor = end
    return spans


def _make_unit_id(story_type: str, section_index: int, container_path: list[str], paragraph_index: int) -> str:
    path_key = "|".join(container_path) if container_path else "root"
    return f"{story_type}:{section_index}:{path_key}:p{paragraph_index}"


def _is_heading_paragraph(paragraph: Paragraph) -> bool:
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").lower()
    return style_name.startswith("heading")


def resolve_unit_paragraph(doc: DocxDocument, unit: DocUnit) -> Optional[Paragraph]:
    if unit.story_type == "body":
        container = doc
    elif unit.story_type == "header":
        if unit.section_index >= len(doc.sections):
            return None
        container = doc.sections[unit.section_index].header
    elif unit.story_type == "footer":
        if unit.section_index >= len(doc.sections):
            return None
        container = doc.sections[unit.section_index].footer
    else:
        return None

    current = container
    i = 0
    while i < len(unit.container_path):
        table_token = unit.container_path[i]
        if not table_token.startswith("table:"):
            return None
        table = _find_table_by_index(current, int(table_token.split(":", 1)[1]))
        if table is None:
            return None
        if i + 2 >= len(unit.container_path):
            return None
        row_token = unit.container_path[i + 1]
        cell_token = unit.container_path[i + 2]
        if not row_token.startswith("row:") or not cell_token.startswith("cell:"):
            return None
        row_index = int(row_token.split(":", 1)[1])
        cell_index = int(cell_token.split(":", 1)[1])
        if row_index >= len(table.rows) or cell_index >= len(table.rows[row_index].cells):
            return None
        current = table.rows[row_index].cells[cell_index]
        i += 3

    return _find_paragraph_by_index(current, unit.paragraph_index)


def _find_table_by_index(container, table_index: int) -> Optional[Table]:
    current_index = 0
    for block_type, block in iter_block_items(container):
        if block_type != "table":
            continue
        if current_index == table_index:
            return block
        current_index += 1
    return None


def _find_paragraph_by_index(container, paragraph_index: int) -> Optional[Paragraph]:
    current_index = 0
    for block_type, block in iter_block_items(container):
        if block_type != "paragraph":
            continue
        if current_index == paragraph_index:
            return block
        current_index += 1
    return None


def reconstruct_docx_from_blueprint(data: DocExportData) -> Optional[io.BytesIO]:
    blueprint = load_blueprint(data.doc_id)
    source_path = get_source_doc_path(data.doc_id)
    if blueprint is None or not source_path or not os.path.exists(source_path):
        return None

    doc = Document(source_path)
    translations_by_unit = _group_translations_by_unit(data.translations)
    reconstructed_count = 0
    unresolved_count = 0

    for unit in blueprint.units:
        paragraph = resolve_unit_paragraph(doc, unit)
        if paragraph is None:
            unresolved_count += 1
            continue

        unit_items = translations_by_unit.get(unit.unit_id)
        if not unit_items:
            continue

        translated_text = "".join(item.translation for item in unit_items if item.translation)
        if translated_text == "":
            continue

        _rewrite_paragraph(paragraph, translated_text, unit.run_spans)
        reconstructed_count += 1

    print(
        f"[docx-blueprint] export doc_id={data.doc_id} "
        f"units={len(blueprint.units)} reconstructed={reconstructed_count} unresolved={unresolved_count}"
    )

    _append_export_footer(doc, data.source_lang, data.target_lang)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _group_translations_by_unit(translations: list[TranslationItem]) -> dict[str, list[TranslationItem]]:
    grouped: dict[str, list[TranslationItem]] = {}
    for item in translations:
        if not item.unit_id:
            continue
        grouped.setdefault(item.unit_id, []).append(item)

    for unit_id, items in grouped.items():
        items.sort(key=lambda value: value.segment_index_within_unit or 0)
    return grouped


def _rewrite_paragraph(paragraph: Paragraph, translated_text: str, run_spans: list[RunSpan]) -> None:
    paragraph.clear()
    if not run_spans:
        paragraph.add_run(translated_text)
        return

    parts = _split_text_across_spans(translated_text, run_spans)
    wrote_any = False
    for span, part in zip(run_spans, parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        _apply_run_span_format(run, span)
        wrote_any = True

    if not wrote_any:
        paragraph.add_run(translated_text)


def _split_text_across_spans(text: str, spans: list[RunSpan]) -> list[str]:
    if len(spans) == 1:
        return [text]

    weights = [max(span.end - span.start, len(span.text), 1) for span in spans]
    remaining_weight = sum(weights)
    remaining_chars = len(text)
    cursor = 0
    pieces: list[str] = []

    for index, weight in enumerate(weights):
        if index == len(weights) - 1:
            pieces.append(text[cursor:])
            break

        if remaining_weight <= 0:
            take = 0
        else:
            take = round(remaining_chars * (weight / remaining_weight))

        if remaining_chars > 0 and take <= 0:
            take = 1
        if take > remaining_chars:
            take = remaining_chars

        pieces.append(text[cursor:cursor + take])
        cursor += take
        remaining_chars -= take
        remaining_weight -= weight

    while len(pieces) < len(spans):
        pieces.append("")
    return pieces


def _apply_run_span_format(run: Run, span: RunSpan) -> None:
    run.bold = span.bold
    run.italic = span.italic
    run.underline = span.underline
    run.font.strike = span.strike
    run.font.all_caps = span.all_caps
    run.font.small_caps = span.small_caps
    if span.font_name:
        run.font.name = span.font_name
    if span.font_size_pt is not None:
        run.font.size = Pt(span.font_size_pt)
    if span.color_rgb:
        try:
            run.font.color.rgb = RGBColor.from_string(span.color_rgb)
        except ValueError:
            pass


def _append_export_footer(doc: DocxDocument, source_lang: str, target_lang: str) -> None:
    doc.add_paragraph()
    divider = doc.add_paragraph()
    p_pr = divider._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    borders.append(bottom)
    p_pr.append(borders)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        f"Translated by TransSync AI  ·  "
        f"{source_lang.upper()} → {target_lang.upper()}  ·  "
        f"{datetime.utcnow().strftime('%B %d, %Y')}"
    )
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)


def make_output_filename(filename: str, target_lang: str) -> str:
    base = filename.rsplit(".", 1)[0] if "." in filename else filename
    return f"translated_{base}_{target_lang}.docx"

