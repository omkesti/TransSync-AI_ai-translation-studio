"""
export.py
---------
Route: POST /api/export

Responsibility:
    Receives the full list of translated sentences + the original raw_text
    from the frontend, reconstructs the document paragraph by paragraph
    (replacing source sentences with their translations), and streams
    a clean DOCX back as a file download.

    Output mirrors the original document structure:
        - Same number of paragraphs as the original
        - Same paragraph order
        - All text translated in-place
        - A single-line attribution footer at the end
        - No title page, no appendix comparison table

Output file format: .docx (Microsoft Word Open XML)
"""

import io
import re
from datetime import datetime
from typing import List, Optional

from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

router = APIRouter()


# ── Request schema ────────────────────────────────────────────────────────────

class TranslationItem(BaseModel):
    source:      str
    translation: str
    match_type:  Optional[str] = None


class ExportRequest(BaseModel):
    doc_id:      str
    filename:    str = "document"
    source_lang: str = "en"
    target_lang: str
    raw_text:    str = ""        # original extracted text with \n\n paragraph separators
    translations: List[TranslationItem]


# ── Heading detection heuristic ───────────────────────────────────────────────

def _is_heading(text: str) -> bool:
    """
    Heuristically detect heading-like paragraphs.
    A paragraph is treated as a heading if it:
      - Is short (≤ 80 chars)
      - Does not end with sentence-ending punctuation
      - Is not just a number
    """
    text = text.strip()
    if not text or len(text) > 80:
        return False
    if text[-1] in ".!?,;:":
        return False
    if re.match(r"^\d+\.?$", text):
        return False
    return True


# ── Paragraph reconstruction ──────────────────────────────────────────────────

def _build_fuzzy_pattern(source: str) -> re.Pattern:
    """
    Builds a case-insensitive regex that matches the source sentence
    while tolerating whitespace variations (e.g., extra spaces or line breaks).
    """
    # Split by any whitespace, escape each word, and join with the regex \s+ token
    words = source.strip().split()
    pattern_str = r"\s+".join(re.escape(word) for word in words)
    return re.compile(pattern_str, flags=re.IGNORECASE)


def _reconstruct_paragraphs(
    raw_text: str,
    translations: List[TranslationItem],
) -> List[str]:
    """
    Splits raw_text into paragraphs (by double newline) and replaces each
    source sentence with its translation using exact substring matching.

    Returns a list of translated paragraph strings in original order.

    Falls back to flat sentence list if raw_text is empty.
    """
    # Build lookup: source text → translated text
    lookup = {item.source.strip(): item.translation for item in translations}

    if not raw_text.strip():
        # Fallback: just return translated sentences one per paragraph
        return [item.translation for item in translations]

    # Split into paragraphs (parser uses \n\n as separator)
    paragraphs = [p.strip() for p in re.split(r"\n{2,}", raw_text) if p.strip()]

    result = []
    for paragraph in paragraphs:
        translated = paragraph
        for source, translation in lookup.items():
            if not source:
                continue
            # Replace only the first occurrence and tolerate whitespace differences.
            pattern = _build_fuzzy_pattern(source)
            translated = pattern.sub(translation, translated, count=1)
        result.append(translated)

    return result


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _add_horizontal_rule(doc: Document) -> None:
    """Inserts a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _reconstruct_docx(req: ExportRequest) -> io.BytesIO:
    """
    Builds the output DOCX:
        1. Translated paragraphs in original document order
        2. A thin horizontal rule separator
        3. One-line italic attribution footer
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    # ── Reconstruct paragraphs ────────────────────────────────────────────────
    translated_paragraphs = _reconstruct_paragraphs(req.raw_text, req.translations)

    for text in translated_paragraphs:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.size = Pt(11)

        if _is_heading(text):
            run.bold = True
            run.font.size = Pt(13)

    # ── Footer separator ──────────────────────────────────────────────────────
    doc.add_paragraph()  # spacer before rule
    _add_horizontal_rule(doc)

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run(
        f"Translated by TransSync AI  ·  "
        f"{req.source_lang.upper()} → {req.target_lang.upper()}  ·  "
        f"{datetime.utcnow().strftime('%B %d, %Y')}"
    )
    footer_run.italic    = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── Serialize ─────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ── Route ─────────────────────────────────────────────────────────────────────

@router.post("/export")
async def export_document(body: ExportRequest):
    """
    POST /api/export

    Accepts the original raw_text and the list of approved translations,
    reconstructs the document paragraph by paragraph with translated text,
    and returns a downloadable DOCX file.

    Request body:
        {
            "doc_id":      "uuid...",
            "filename":    "annual_report.pdf",
            "source_lang": "en",
            "target_lang": "fr",
            "raw_text":    "Full original extracted text with \\n\\n paragraph breaks...",
            "translations": [
                { "source": "...", "translation": "...", "match_type": "tm_exact" },
                ...
            ]
        }

    Response:
        Binary DOCX stream (Content-Disposition: attachment)
        Structure: translated paragraphs in original order + attribution footer
    """
    docx_buf = _reconstruct_docx(body)

    # Derive a clean output filename
    base = body.filename.rsplit(".", 1)[0] if "." in body.filename else body.filename
    out_name = f"translated_{base}_{body.target_lang}.docx"

    return StreamingResponse(
        docx_buf,
        media_type=(
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ),
        headers={
            "Content-Disposition": f'attachment; filename="{out_name}"'
        },
    )
